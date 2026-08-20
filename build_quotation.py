from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import date, timedelta
from pathlib import Path


OUT = Path("Quotation_Vedanta_Mero_Attendance_Customization.docx")
TODAY = date(2026, 8, 19)
VALID_UNTIL = TODAY + timedelta(days=15)

NAVY = "17324D"
TEAL = "0D7C78"
LIGHT_TEAL = "EAF6F5"
PALE_BLUE = "EEF3F7"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "66717C"
DARK = "1F2933"
WHITE = "FFFFFF"
BORDER = "CAD3DC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in edges:
            continue
        edge_data = edges[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, bold=False, color=DARK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_cell_text(cell, size=9.5, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        for run in p.runs:
            set_run(run, size=size, bold=bold, color=color)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run(run, size=9, color=MID_GRAY)


def add_real_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] + [-1]) + 1
    num_id = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] + [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def bullet(doc, text, num_id):
    p = doc.add_paragraph(style="Normal")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def label_value_paragraph(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label + " ")
    set_run(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(value)
    set_run(r, size=10.5)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, NAVY, 16, 8),
    ("Heading 2", 13, TEAL, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Quote Note" not in styles:
    quote_note = styles.add_style("Quote Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    quote_note = styles["Quote Note"]
quote_note.font.name = "Calibri"
quote_note._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
quote_note._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
quote_note.font.size = Pt(9.5)
quote_note.font.color.rgb = RGBColor.from_string(MID_GRAY)
quote_note.paragraph_format.space_before = Pt(4)
quote_note.paragraph_format.space_after = Pt(4)

num_id = add_real_bullet_numbering(doc)

# Quiet running header and footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
hr = hp.add_run("WAKE AND TECH NEPAL  |  COMMERCIAL QUOTATION")
set_run(hr, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)
fr = fp.add_run("Wake and Tech Nepal  •  Confidential  •  Page ")
set_run(fr, size=8.5, color=MID_GRAY)
add_field(fp, "PAGE")

# Customer-pack opening block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("SOFTWARE CUSTOMIZATION QUOTATION")
set_run(r, size=10, bold=True, color=TEAL)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Vedanta Attendance & Workforce App")
set_run(r, size=25, bold=True, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Customization of Mero Attendance into a Vedanta-branded iOS, Android and web application")
set_run(r, size=12.5, color=MID_GRAY)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(meta, [4680, 4680], indent=120)
meta_data = [
    ("Prepared for", "Vedanta Publication Private Limited"),
    ("Prepared by", "Wake and Tech Nepal"),
    ("Quotation no.", "WTN-QTN-2026-0819-01"),
    ("Date / Valid until", f"{TODAY.strftime('%d %B %Y')}  /  {VALID_UNTIL.strftime('%d %B %Y')}"),
]
for i, (label, value) in enumerate(meta_data):
    row = meta.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for j, cell in enumerate(row.cells):
        set_cell_margins(cell, top=75, bottom=75)
        set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": BORDER})
        style_cell_text(cell, size=9.5, bold=(j == 0), color=(TEAL if j == 0 else DARK))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Commercial summary")
set_run(r, size=16, bold=True, color=NAVY)

pricing = doc.add_table(rows=1, cols=4)
pricing.alignment = WD_TABLE_ALIGNMENT.LEFT
pricing.rows[0].cells[0].text = "Package"
pricing.rows[0].cells[1].text = "Included modules"
pricing.rows[0].cells[2].text = "One-time customization"
pricing.rows[0].cells[3].text = "Annual AMS + hosting"
rows = [
    ("Option A\nCore Workforce", "Attendance, payroll and leave", "NPR 300,000", "NPR 50,000 / year"),
    ("Option B\nIntegrated Operations", "Everything in Option A, plus task management, stock management and financial management", "NPR 500,000", "NPR 60,000 / year"),
]
for data in rows:
    cells = pricing.add_row().cells
    for cell, value in zip(cells, data):
        cell.text = value
set_table_geometry(pricing, [1750, 3510, 2050, 2050], indent=120)
set_repeat_table_header(pricing.rows[0])
for i, row in enumerate(pricing.rows):
    for j, cell in enumerate(row.cells):
        set_cell_margins(cell, top=115, bottom=115)
        set_cell_border(cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            start={"val": "single", "sz": "6", "color": BORDER},
            end={"val": "single", "sz": "6", "color": BORDER})
        if i == 0:
            set_cell_shading(cell, NAVY)
            style_cell_text(cell, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            if i == 2:
                set_cell_shading(cell, LIGHT_TEAL)
            style_cell_text(cell, size=9.2, bold=(j in (0, 2, 3)), color=(NAVY if j in (0, 2, 3) else DARK), align=(WD_ALIGN_PARAGRAPH.CENTER if j in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT))

p = doc.add_paragraph(style="Quote Note")
r = p.add_run("Prices are in Nepalese Rupees (NPR). Applicable taxes, if any, will be added as required by law. Option B is a complete package price and is not added on top of Option A.")
set_run(r, size=9.5, italic=True, color=MID_GRAY)

doc.add_heading("Scope and deliverables", level=1)
bullet(doc, "A dedicated Vedanta-branded mobile application for iOS and Android, customized from the agreed Mero Attendance feature set.", num_id)
bullet(doc, "A dedicated web application accessible at vedanta.meroattendance.com.", num_id)
bullet(doc, "Configuration and customization of the modules included in the selected package.", num_id)
bullet(doc, "Initial deployment support and transition to annual application management services after go-live.", num_id)

doc.add_heading("Platform ownership and management", level=1)
label_value_paragraph("Product ownership.", "Vedanta will receive and use its own branded application with the features mutually agreed for the selected package.")
label_value_paragraph("Source-code ownership.", "The application source code, underlying platform, reusable components, frameworks and associated intellectual property will remain the exclusive property of Wake and Tech Nepal. No source-code transfer or assignment is included in this quotation.")
label_value_paragraph("Application management.", "Wake and Tech Nepal will manage the application and provide issue resolution, corrective maintenance and bug fixes under the applicable annual AMS plan.")

doc.add_page_break()

doc.add_heading("Annual AMS and hosting", level=1)
p = doc.add_paragraph()
r = p.add_run("The annual Application Management Service (AMS) fee is billed yearly from the service commencement date agreed at go-live. It includes:")
set_run(r, size=10.5)
bullet(doc, "Hosting for the Vedanta web application and supporting production environment within Wake and Tech Nepal's managed infrastructure.", num_id)
bullet(doc, "Routine monitoring, application management, issue investigation, bug fixes and corrective maintenance for the agreed features.", num_id)
bullet(doc, "Reasonable deployment support for maintenance releases affecting the agreed scope.", num_id)

doc.add_heading("Assumptions and exclusions", level=1)
bullet(doc, "Detailed workflows, reports, user roles, branding assets and acceptance criteria will be confirmed in a mutually approved scope or work order before development starts.", num_id)
bullet(doc, "New modules, new features, major workflow changes, integrations, data migration or functionality beyond the agreed package are not included and will be quoted separately.", num_id)
bullet(doc, "Apple App Store, Google Play, SMS, email, payment-gateway or other third-party account and usage charges are excluded unless specifically included in the final work order.", num_id)
bullet(doc, "Vedanta will provide required logos, brand guidelines, content, authorized contacts and timely approvals.", num_id)
bullet(doc, "Service availability depends on third-party infrastructure and platform policies. Scheduled maintenance and events outside Wake and Tech Nepal's reasonable control are excluded from defect liability.", num_id)

doc.add_heading("Commercial terms", level=1)
terms = doc.add_table(rows=6, cols=2)
terms_data = [
    ("Quotation validity", f"15 days, through {VALID_UNTIL.strftime('%d %B %Y')}"),
    ("Package selection", "Vedanta may select either Option A or Option B."),
    ("Payment schedule", "To be agreed in the signed work order or service agreement."),
    ("Implementation timeline", "To be finalized after scope confirmation and receipt of required inputs."),
    ("Acceptance", "Deliverables will be accepted against the mutually agreed scope and acceptance criteria."),
    ("Annual renewal", "AMS and hosting renew annually at the stated package rate, subject to any mutually agreed changes in scope or third-party costs."),
]
for i, (label, value) in enumerate(terms_data):
    terms.rows[i].cells[0].text = label
    terms.rows[i].cells[1].text = value
set_table_geometry(terms, [2500, 6860], indent=120)
for i, row in enumerate(terms.rows):
    for j, cell in enumerate(row.cells):
        set_cell_margins(cell, top=95, bottom=95)
        set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": BORDER})
        if j == 0:
            set_cell_shading(cell, PALE_BLUE)
        style_cell_text(cell, size=9.5, bold=(j == 0), color=(NAVY if j == 0 else DARK))

doc.add_page_break()
doc.add_heading("Acceptance", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("This quotation becomes binding only when incorporated into a mutually signed work order or service agreement. The selected package should be marked below.")
set_run(r, size=10.5)

choice = doc.add_table(rows=1, cols=2)
choice.rows[0].cells[0].text = "☐  Option A — Core Workforce"
choice.rows[0].cells[1].text = "☐  Option B — Integrated Operations"
set_table_geometry(choice, [4680, 4680], indent=120)
for cell in choice.rows[0].cells:
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_margins(cell, top=130, bottom=130)
    set_cell_border(cell,
        top={"val": "single", "sz": "8", "color": TEAL},
        bottom={"val": "single", "sz": "8", "color": TEAL},
        start={"val": "single", "sz": "8", "color": TEAL},
        end={"val": "single", "sz": "8", "color": TEAL})
    style_cell_text(cell, size=10, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
sign = doc.add_table(rows=4, cols=2)
sign_data = [
    ("For Vedanta Publication Private Limited", "For Wake and Tech Nepal"),
    ("Authorized signature: ____________________", "Authorized signature: ____________________"),
    ("Name / title: ___________________________", "Name / title: ___________________________"),
    ("Date: __________________________________", "Date: __________________________________"),
]
for i, rowdata in enumerate(sign_data):
    for j, value in enumerate(rowdata):
        sign.rows[i].cells[j].text = value
set_table_geometry(sign, [4680, 4680], indent=120)
for i, row in enumerate(sign.rows):
    for cell in row.cells:
        set_cell_margins(cell, top=(110 if i == 0 else 85), bottom=(110 if i == 0 else 85))
        if i == 0:
            set_cell_shading(cell, NAVY)
            style_cell_text(cell, size=9.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            style_cell_text(cell, size=9.5, color=DARK)

# Document properties and save.
doc.core_properties.title = "Quotation - Vedanta Attendance & Workforce App"
doc.core_properties.subject = "Customization of Mero Attendance for Vedanta Publication Private Limited"
doc.core_properties.author = "Wake and Tech Nepal"
doc.core_properties.keywords = "quotation, Vedanta, Mero Attendance, application customization, AMS"
doc.core_properties.comments = "Prepared for Vedanta Publication Private Limited"
doc.save(OUT)
print(OUT.resolve())
