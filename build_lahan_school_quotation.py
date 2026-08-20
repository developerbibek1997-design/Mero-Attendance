from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import build_quotation as base


OUT = Path("Quotation_Lahan_School_Mero_Attendance_App.docx")
TODAY = date(2026, 8, 19)
VALID_UNTIL = TODAY + timedelta(days=15)

NAVY = base.NAVY
TEAL = base.TEAL
LIGHT_TEAL = base.LIGHT_TEAL
PALE_BLUE = base.PALE_BLUE
MID_GRAY = base.MID_GRAY
DARK = base.DARK
WHITE = base.WHITE
BORDER = base.BORDER


def add_text(doc, text, size=10.5, bold=False, color=DARK, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    base.set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def shade_and_style_table(table, header=True, highlight_row=None, numeric_cols=()):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            base.set_cell_margins(cell, top=105, bottom=105, start=120, end=120)
            base.set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": BORDER},
                bottom={"val": "single", "sz": "6", "color": BORDER},
                start={"val": "single", "sz": "6", "color": BORDER},
                end={"val": "single", "sz": "6", "color": BORDER},
            )
            align = WD_ALIGN_PARAGRAPH.CENTER if j in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            if header and i == 0:
                base.set_cell_shading(cell, NAVY)
                base.style_cell_text(cell, size=8.8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if highlight_row is not None and i == highlight_row:
                    base.set_cell_shading(cell, LIGHT_TEAL)
                base.style_cell_text(cell, size=9.2, bold=(j in numeric_cols), color=(NAVY if j in numeric_cols else DARK), align=align)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, NAVY, 14, 7),
    ("Heading 2", 13, TEAL, 10, 5),
    ("Heading 3", 12, NAVY, 8, 4),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Quote Note" not in styles:
    note = styles.add_style("Quote Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    note = styles["Quote Note"]
note.font.name = "Calibri"
note.font.size = Pt(9.5)
note.font.color.rgb = RGBColor.from_string(MID_GRAY)
note.paragraph_format.space_before = Pt(4)
note.paragraph_format.space_after = Pt(4)

num_id = base.add_real_bullet_numbering(doc)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.paragraph_format.space_after = Pt(0)
base.set_run(hp.add_run("WAKE AND TECH NEPAL  |  COMMERCIAL QUOTATION"), size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_after = Pt(0)
base.set_run(fp.add_run("Wake and Tech Nepal  •  Confidential  •  Page "), size=8.5, color=MID_GRAY)
base.add_field(fp, "PAGE")

# Customer-pack opening block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
base.set_run(p.add_run("MERO ATTENDANCE APP CUSTOMIZATION QUOTATION"), size=10, bold=True, color=TEAL)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
base.set_run(p.add_run("Lahan School Mero Attendance App"), size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
base.set_run(
    p.add_run("Customization of Mero Attendance into a Lahan School-branded iOS, Android and web app with student, library and billing management"),
    size=12.2,
    color=MID_GRAY,
)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
base.set_table_geometry(meta, [4680, 4680], indent=120)
meta_data = [
    ("Prepared for", "Lahan School"),
    ("Prepared by", "Wake and Tech Nepal"),
    ("Quotation no.", "WTN-QTN-2026-0819-02"),
    ("Date / Valid until", f"{TODAY.strftime('%d %B %Y')}  /  {VALID_UNTIL.strftime('%d %B %Y')}"),
]
for i, (label, value) in enumerate(meta_data):
    meta.rows[i].cells[0].text = label
    meta.rows[i].cells[1].text = value
    for j, cell in enumerate(meta.rows[i].cells):
        base.set_cell_margins(cell, top=70, bottom=70)
        base.set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": BORDER})
        base.style_cell_text(cell, size=9.5, bold=(j == 0), color=(TEAL if j == 0 else DARK))

doc.add_heading("Commercial summary", level=1)
pricing = doc.add_table(rows=1, cols=4)
pricing.alignment = WD_TABLE_ALIGNMENT.LEFT
for cell, text in zip(pricing.rows[0].cells, ("Licensed capacity", "Included systems", "One-time implementation", "Annual AMS + hosting")):
    cell.text = text
pricing_rows = [
    ("Up to 500 users", "Customized Mero Attendance app with student, library and billing", "NPR 300,000", "NPR 50,000 / year"),
    ("Up to 1,000 users", "Customized Mero Attendance app with student, library and billing", "NPR 400,000", "NPR 55,000 / year"),
    ("Up to 2,000 users", "Customized Mero Attendance app with student, library and billing", "NPR 500,000", "NPR 60,000 / year"),
]
for data in pricing_rows:
    cells = pricing.add_row().cells
    for cell, value in zip(cells, data):
        cell.text = value
base.set_table_geometry(pricing, [1750, 3260, 2175, 2175], indent=120)
base.set_repeat_table_header(pricing.rows[0])
shade_and_style_table(pricing, header=True, highlight_row=2, numeric_cols=(2, 3))

p = doc.add_paragraph(style="Quote Note")
base.set_run(
    p.add_run("Prices are in Nepalese Rupees (NPR). Applicable taxes, if any, will be added as required by law. Capacity refers to active student and staff accounts combined."),
    size=9.4,
    italic=True,
    color=MID_GRAY,
)

doc.add_heading("Mero Attendance customization scope", level=1)
base.bullet(doc, "A dedicated Lahan School-branded mobile application for iOS and Android, customized from the agreed Mero Attendance feature set.", num_id)
base.bullet(doc, "A dedicated web application accessible at lahan.meroattendance.com.", num_id)
base.bullet(doc, "Student management: admissions and profiles, class/section records, guardian details, enrollment status and standard student reporting.", num_id)
base.bullet(doc, "Library management: book catalogue, member records, issue/return tracking, due dates, availability and basic circulation reports.", num_id)
base.bullet(doc, "Billing and fee management: fee structures, student invoices, payment entry, outstanding balances, receipts and standard collection reports.", num_id)

doc.add_page_break()
doc.add_heading("Usage-based add-ons", level=1)
addons = doc.add_table(rows=1, cols=4)
for cell, text in zip(addons.rows[0].cells, ("Add-on", "Unit rate", "Illustrative quantity", "Illustrative total")):
    cell.text = text
addon_rows = [
    ("RFID ID card - normal quality", "NPR 120 / card", "500 / 1,000 / 2,000 cards", "NPR 60,000 / 120,000 / 240,000"),
    ("RFID ID card - better quality", "NPR 150 / card", "500 / 1,000 / 2,000 cards", "NPR 75,000 / 150,000 / 300,000"),
    ("SMS integration usage", "NPR 0.50 / SMS", "Actual messages sent", "Quantity × NPR 0.50"),
]
for data in addon_rows:
    cells = addons.add_row().cells
    for cell, value in zip(cells, data):
        cell.text = value
base.set_table_geometry(addons, [2700, 1700, 2250, 2710], indent=120)
base.set_repeat_table_header(addons.rows[0])
shade_and_style_table(addons, header=True, highlight_row=2, numeric_cols=(1,))

p = doc.add_paragraph(style="Quote Note")
base.set_run(
    p.add_run("RFID card totals are illustrations only; billing is based on the actual approved quantity. Replacement and reprint cards are charged at the same applicable unit rate. SMS charges are usage-based."),
    size=9.4,
    italic=True,
    color=MID_GRAY,
)

doc.add_heading("Annual AMS and hosting", level=1)
add_text(doc, "The annual Application Management Service (AMS) fee is billed yearly from the service commencement date agreed at go-live. It includes:")
base.bullet(doc, "Managed hosting for the Lahan School web app and the supporting Mero Attendance production environment.", num_id)
base.bullet(doc, "Application management, routine monitoring, issue investigation, bug fixes and corrective maintenance for the agreed features.", num_id)
base.bullet(doc, "Reasonable deployment support for maintenance releases affecting the agreed Mero Attendance customization scope.", num_id)

doc.add_heading("Platform ownership and commercial assumptions", level=1)
base.bullet(doc, "Lahan School will receive and use its own branded application with the Mero Attendance features and school-management modules mutually agreed in the signed scope.", num_id)
base.bullet(doc, "The application source code, underlying Mero Attendance platform, reusable components and associated intellectual property remain the exclusive property of Wake and Tech Nepal; no source-code transfer is included.", num_id)
base.bullet(doc, "New modules, major workflow changes, data migration, custom integrations, payment gateways, biometric/RFID readers, printers and other hardware are excluded unless specifically quoted.", num_id)
base.bullet(doc, "SMS delivery is subject to supported telecom routes and provider availability; the stated rate applies per successfully submitted SMS and may be revised if the upstream provider changes its tariff.", num_id)

doc.add_page_break()
doc.add_heading("Commercial terms", level=1)
terms = doc.add_table(rows=6, cols=2)
term_rows = [
    ("Quotation validity", f"15 days, through {VALID_UNTIL.strftime('%d %B %Y')}"),
    ("Capacity selection", "Lahan School may select the 500-, 1,000- or 2,000-user tier."),
    ("Payment schedule", "To be agreed in the signed work order or service agreement."),
    ("Implementation timeline", "To be finalized after scope confirmation and receipt of required school data and approvals."),
    ("Acceptance", "Deliverables will be accepted against mutually agreed scope and acceptance criteria."),
    ("Annual renewal", "AMS and hosting renew annually at the selected tier rate, subject to mutually agreed scope or third-party cost changes."),
]
for i, (label, value) in enumerate(term_rows):
    terms.rows[i].cells[0].text = label
    terms.rows[i].cells[1].text = value
base.set_table_geometry(terms, [2500, 6860], indent=120)
for row in terms.rows:
    for j, cell in enumerate(row.cells):
        base.set_cell_margins(cell, top=85, bottom=85)
        base.set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": BORDER})
        if j == 0:
            base.set_cell_shading(cell, PALE_BLUE)
        base.style_cell_text(cell, size=9.3, bold=(j == 0), color=(NAVY if j == 0 else DARK))

doc.add_heading("Acceptance", level=1)
add_text(doc, "This quotation becomes binding only when incorporated into a mutually signed work order or service agreement. Mark the selected capacity below.", after=8)

choice = doc.add_table(rows=1, cols=3)
for cell, value in zip(choice.rows[0].cells, ("☐  Up to 500 users", "☐  Up to 1,000 users", "☐  Up to 2,000 users")):
    cell.text = value
base.set_table_geometry(choice, [3120, 3120, 3120], indent=120)
for cell in choice.rows[0].cells:
    base.set_cell_shading(cell, LIGHT_TEAL)
    base.set_cell_margins(cell, top=120, bottom=120)
    base.set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": TEAL},
        bottom={"val": "single", "sz": "8", "color": TEAL},
        start={"val": "single", "sz": "8", "color": TEAL},
        end={"val": "single", "sz": "8", "color": TEAL},
    )
    base.style_cell_text(cell, size=9.3, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(0)
sign = doc.add_table(rows=4, cols=2)
sign_rows = [
    ("For Lahan School", "For Wake and Tech Nepal"),
    ("Authorized signature: ____________________", "Authorized signature: ____________________"),
    ("Name / title: ___________________________", "Name / title: ___________________________"),
    ("Date: __________________________________", "Date: __________________________________"),
]
for i, data in enumerate(sign_rows):
    sign.rows[i].cells[0].text = data[0]
    sign.rows[i].cells[1].text = data[1]
base.set_table_geometry(sign, [4680, 4680], indent=120)
for i, row in enumerate(sign.rows):
    for cell in row.cells:
        base.set_cell_margins(cell, top=(95 if i == 0 else 70), bottom=(95 if i == 0 else 70))
        if i == 0:
            base.set_cell_shading(cell, NAVY)
            base.style_cell_text(cell, size=9.4, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            base.style_cell_text(cell, size=9.2, color=DARK)

doc.core_properties.title = "Quotation - Lahan School Mero Attendance App"
doc.core_properties.subject = "Customization of Mero Attendance for Lahan School with student, library and billing management"
doc.core_properties.author = "Wake and Tech Nepal"
doc.core_properties.keywords = "Lahan School, Mero Attendance, student management, library management, billing, RFID, SMS"
doc.core_properties.comments = "Prepared for Lahan School"
doc.save(OUT)
print(OUT.resolve())
